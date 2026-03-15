import io
import json
import logging
import os
import sys
from typing import Dict, Any, Generator, Tuple, List

import docx
import openpyxl
import pypdf
from langchain_core.documents import Document

from DataLoaderS3Service import DataLoaderS3Service
from buissnes_agent.MetadataModels import FileMetadata

logging.basicConfig(level=logging.INFO, stream=sys.stderr)
logger = logging.getLogger(__name__)

class DataLoaderS3FileLoader:
    def __init__(self, bucket_name: str, prefix: str):
        self.bucket_name = bucket_name

        # 1. Usuwamy białe znaki
        clean_prefix = prefix.strip() if prefix else ""

        # 2. Jeśli prefix jest podany, upewniamy się, że kończy się slashem "/"
        if clean_prefix:
            if not clean_prefix.endswith('/'):
                self.prefix = f"{clean_prefix}/"
            else:
                self.prefix = clean_prefix
            logger.info(f"S3FileLoader: Ustawiono filtr na folder: '{self.prefix}'")
        else:
            # 3. Jeśli prefix jest pusty -> OSTRZEŻENIE
            self.prefix = ""
            logger.warning("!!! UWAGA: Nie podano folderu (prefix). Skrypt pobierze CAŁY BUCKET !!!")

        self.s3_service = DataLoaderS3Service()

    def list_objects(self) -> Generator[str, None, None]:
        # Przekazujemy prefix do serwisu S3. AWS zwróci tylko obiekty zaczynające się od tego ciągu.
        return self.s3_service.list_objects(self.bucket_name, self.prefix)

    def load_file_with_metadata(self, s3_key: str) -> Tuple[List[Document], Dict[str, Any]]:
        filename = os.path.basename(s3_key)
        ext = os.path.splitext(s3_key)[1].lower()

        domain_value = self._extract_domain_first(s3_key)

        # Tworzenie metadanych
        meta_obj = FileMetadata(
            source=f"s3://{self.bucket_name}/{s3_key}",
            title=filename,
            extension=ext,
            url=f"https://{self.bucket_name}.s3.amazonaws.com/{s3_key}",
            domain=domain_value,  # <--- Podstawienie dynamicznej domeny

            # Dodajemy domenę do tagów (z małych liter) dla lepszego filtrowania w bazie wektorowej
            tags=["s3_storage", "cloud", domain_value.lower()],

            page_number=None  # Cały plik, więc brak konkretnej strony
        )

        documents: List[Document] = []
        base_metadata = meta_obj.to_dict()

        try:
            # --- ZMIANA: Pobieramy bajty do pamięci RAM (bez zapisu na dysk) ---
            file_bytes = self.s3_service.download_bytes(self.bucket_name, s3_key)
            file_stream = io.BytesIO(file_bytes)

            # --- XLSX (Excel) ---
            if ext == ".xlsx":
                documents = self._process_xlsx(file_stream, filename)
            # --- PDF ---
            elif ext == ".pdf":
                documents = self._process_pdf(file_stream, filename)
            # --- DOCX ---
            elif ext == ".docx":
                documents = self._process_docx(file_stream, filename)
            # --- TEKSTOWE (TXT, MD, XML, JSON) ---
            else:
                text_content = file_bytes.decode("utf-8", errors="ignore")
                if text_content.strip():
                    documents = [Document(page_content=text_content, metadata={"page_number": 1})]

            # Wzbogacamy każdą stronę (Document) o globalne metadane pliku S3
            for doc in documents:
                merged_meta = base_metadata.copy()
                merged_meta.update(doc.metadata)
                doc.metadata = merged_meta

            return documents, base_metadata

        except Exception as e:
            logger.error(f"Krytyczny błąd przy pliku {s3_key}: {e}")
            return [], {}  # <--- POPRAWKA: Zwraca pustą listę zamiast "", zapobiega błędom w Orkiestratorze

    def _process_xlsx(self, file_stream: io.BytesIO, filename: str) -> List[Document]:
        """
        Pobiera plik XLSX z S3 i konwertuje go na JSON String.
        """
        documents = []
        try:
            # openpyxl potrafi czytać bezpośrednio ze strumienia BytesIO
            wb = openpyxl.load_workbook(file_stream, data_only=True)

            for i, sheet in enumerate(wb.worksheets):
                sheet_data = []
                for row in sheet.iter_rows(values_only=True):
                    if any(cell is not None for cell in row):
                        clean_row = [str(cell) if cell is not None else None for cell in row]
                        sheet_data.append(clean_row)

                if sheet_data:
                    sheet_json_str = json.dumps({sheet.title: sheet_data}, ensure_ascii=False)
                    doc = Document(
                        page_content=sheet_json_str,
                        metadata={"page_number": i + 1, "sheet_name": sheet.title}
                    )
                    documents.append(doc)
            return documents
        except Exception as e:
            logger.error(f"Błąd przetwarzania XLSX {filename}: {e}")
            return []

    def _process_pdf(self, file_stream: io.BytesIO, filename: str) -> List[Document]:
        """
         Pobiera PDF z S3 i ekstrahuje tekst.
         """
        documents = []
        try:
            # pypdf również czyta strumienie z pamięci
            reader = pypdf.PdfReader(file_stream)

            for i, page in enumerate(reader.pages):
                extracted = page.extract_text()
                if extracted and extracted.strip():
                    doc = Document(
                        page_content=extracted,
                        metadata={"page_number": i + 1}
                    )
                    documents.append(doc)
            return documents
        except Exception as e:
            logger.error(f"Błąd przetwarzania PDF {filename}: {e}")
            return []

    def _process_docx(self, file_stream: io.BytesIO, filename: str) -> List[Document]:
        """
        Pobiera DOCX z S3 (strumień RAM) i symuluje podział na strony
        szukając w kodzie XML twardych enterów i miękkich podziałów Worda.
        """
        documents = []
        try:
            doc = docx.Document(file_stream)

            current_page_text = []
            current_char_count = 0
            page_num = 1

            # Zakładamy średnio 2500 znaków na stronę A4 (możesz to dostosować)
            MAX_CHARS_PER_PAGE = 2500

            for para in doc.paragraphs:
                # 1. Sprawdzamy tagi w XML
                hard_breaks = para._element.xpath('.//w:br[@w:type="page"]')
                rendered_breaks = para._element.xpath('.//w:lastRenderedPageBreak')

                # 2. Sprawdzamy heurystykę (czy uzbieraliśmy już tyle tekstu, że to na pewno nowa strona)
                limit_reached = current_char_count >= MAX_CHARS_PER_PAGE

                # Jeśli wystąpił fizyczny znacznik LUB przekroczyliśmy limit znaków
                if hard_breaks or rendered_breaks or limit_reached:

                    joined_text = "\n".join(current_page_text).strip()
                    if joined_text:
                        documents.append(Document(page_content=joined_text, metadata={"page_number": page_num}))

                    # Resetujemy bufory na nową stronę
                    current_page_text = []
                    current_char_count = 0
                    page_num += 1

                # Zbieramy tekst z obecnego paragrafu
                text = para.text.strip()
                if text:
                    current_page_text.append(text)
                    current_char_count += len(text)

            # Zapisujemy to, co zostało po zakończeniu pętli (ostatnia strona)
            if current_page_text:
                joined_text = "\n".join(current_page_text).strip()
                if joined_text:
                    documents.append(Document(page_content=joined_text, metadata={"page_number": page_num}))

            return documents
        except Exception as e:
            logger.error(f"Błąd parsowania DOCX {filename}: {e}")
            return []

    def _process_text(self, file_bytes: bytes, filename: str) -> List[Document]:
        """
        Przetwarza pobrane bajty z S3 (TXT, MD, XML, JSON, etc.) i zwraca 1 stronę.
        """
        try:
            text_content = file_bytes.decode("utf-8", errors="ignore")
            if text_content.strip():
                return [Document(page_content=text_content, metadata={"page_number": 1})]
            return []
        except Exception as e:
            logger.error(f"Błąd dekodowania tekstu {filename}: {e}")
            return []

    def _extract_domain_first(self, s3_key: str) -> str:
        """
        Zawsze pobiera GŁÓWNY (najwyższy) folder wprost z klucza S3.
        Np. dla klucza "technical/ISO20022/MDR/plik.pdf" zwróci "technical".
        """
        # Zabezpieczenie: usuwamy ewentualne ukośniki na samym początku klucza
        clean_key = s3_key.lstrip("/")
        parts = clean_key.split('/')

        if len(parts) > 1:
            # Bierzemy ZAWSZE pierwszy, najwyższy folder w hierarchii bucketa
            return parts[0]
        else:
            # Plik leży bezpośrednio w roocie bucketa, bez żadnego folderu
            return "general"