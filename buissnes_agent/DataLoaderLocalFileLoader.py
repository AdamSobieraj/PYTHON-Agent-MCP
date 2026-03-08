import json
import logging
import os
import sys
from typing import Generator, Tuple, Dict, Any, List

import docx
import openpyxl
import pypdf
from langchain_core.documents import Document

from buissnes_agent.MetadataModels import FileMetadata
from buissnes_agent.config_loader import settings

logging.basicConfig(level=logging.INFO, stream=sys.stderr)
logger = logging.getLogger(__name__)


class DataLoaderLocalFileLoader:
    """
    Adapter dla plików lokalnych.
    Obsługuje: TXT, MD, XML, JSON, PDF, DOCX, XLSX.
    """

    def __init__(self, directory: str):
        self.directory = os.path.abspath(directory)

    def list_objects(self) -> Generator[str, None, None]:

        allowed_exts = settings.get("chunking.allowed_extensions", [])
        ext_tuple = tuple(allowed_exts)

        if not os.path.exists(self.directory):
            logger.error(f"Katalog nie istnieje: {self.directory}")
            return

        for root, _, files in os.walk(self.directory):
            for file in files:
                ext = os.path.splitext(file)[1].lower()
                if ext in ext_tuple:
                    yield os.path.join(root, file)

    def load_file_with_metadata(self, file_path: str) -> Tuple[List[Document], Dict[str, Any]]:
        filename = os.path.basename(file_path)
        ext = os.path.splitext(file_path)[1].lower()

        domain_value = self._extract_domain_first(file_path)

        # 1. Tworzenie obiektu metadanych (Type Safe)
        # page_number zostawiamy domyślnie puste, nadpiszemy je na poziomie poszczególnych Documentów
        meta_obj = FileMetadata(
            source=f"file://{file_path}",
            title=filename,
            extension=ext,
            url=f"file://{file_path}",
            domain=domain_value,
            tags=["local", "filesystem"],
            page_number=None  # Domyślny brak strony
        )

        documents: List[Document] = []
        base_metadata = meta_obj.to_dict()

        try:
            # --- XLSX (Excel) ---
            if ext == ".xlsx":
                documents = self._process_xlsx(file_path)
            # --- PDF ---
            elif ext == ".pdf":
                documents = self._process_pdf(file_path)
            # --- DOCX ---
            elif ext == ".docx":
                documents = self._process_docx(file_path)
            # --- TEKSTOWE ---
            else:
                documents = self._process_text(file_path)

            # --- ZMIANA: Wzbogacamy każdą stronę (Document) o globalne metadane pliku ---
            for doc in documents:
                # Kopiujemy bazowe metadane pliku
                merged_meta = base_metadata.copy()
                # Nadpisujemy je specyficznymi metadanymi ze strony (np. page_number z PDF)
                merged_meta.update(doc.metadata)
                # Zapisujemy połączone metadane z powrotem do dokumentu
                doc.metadata = merged_meta

            return documents, base_metadata

        except Exception as e:
            logger.error(f"Krytyczny błąd przy pliku {file_path}: {e}")
            return [], {}

    def _process_xlsx(self, file_path: str) -> List[Document]:
        """
        Prywatna metoda do konwersji pliku XLSX na obiekty Document.
        Każdy arkusz (sheet) staje się osobną "stroną" (Documentem).
        """
        documents = []
        try:
            wb = openpyxl.load_workbook(file_path, data_only=True)

            for i, sheet in enumerate(wb.worksheets):
                sheet_data = []

                for row in sheet.iter_rows(values_only=True):
                    if any(cell is not None for cell in row):
                        clean_row = [str(cell) if cell is not None else None for cell in row]
                        sheet_data.append(clean_row)

                if sheet_data:
                    sheet_json_str = json.dumps({sheet.title: sheet_data}, ensure_ascii=False)
                    doc = Document(
                        page_content=sheet_json_str,
                        metadata={"page_number": i + 1, "sheet_name": sheet.title}
                    )
                    documents.append(doc)

            if not documents:
                logger.warning(f"Plik XLSX {os.path.basename(file_path)} jest pusty lub nie zawiera danych.")

            return documents

        except Exception as e:
            logger.error(f"Błąd przetwarzania XLSX {os.path.basename(file_path)}: {e}")
            return []

    # --- Zwraca List[Document], traktując każdy arkusz jako "stronę" ---
    def _process_docx(self, file_path: str) -> List[Document]:

        """
        Ulepszona ekstrakcja tekstu z DOCX z symulacją podziału na strony.
        Szuka w kodzie XML znaczników twardego podziału strony (page break)
        oraz wyliczonych podziałów (lastRenderedPageBreak).
        """
        documents = []
        try:
            doc = docx.Document(file_path)

            current_page_text = []
            current_char_count = 0
            page_num = 1

            # Zakładamy średnio 2500 znaków na stronę A4 (możesz to dostosować)
            MAX_CHARS_PER_PAGE = 2500

            for para in doc.paragraphs:
                # 1. Sprawdzamy tagi w XML
                hard_breaks = para._element.xpath('.//w:br[@w:type="page"]')
                rendered_breaks = para._element.xpath('.//w:lastRenderedPageBreak')

                # 2. Sprawdzamy heurystykę (czy uzbieraliśmy już tyle tekstu, że to na pewno nowa strona)
                limit_reached = current_char_count >= MAX_CHARS_PER_PAGE

                # Jeśli wystąpił fizyczny znacznik LUB przekroczyliśmy limit znaków
                if hard_breaks or rendered_breaks or limit_reached:

                    joined_text = "\n".join(current_page_text).strip()
                    if joined_text:
                        documents.append(Document(page_content=joined_text, metadata={"page_number": page_num}))

                    # Resetujemy bufory na nową stronę
                    current_page_text = []
                    current_char_count = 0
                    page_num += 1

                # Zbieramy tekst z obecnego paragrafu
                text = para.text.strip()
                if text:
                    current_page_text.append(text)
                    current_char_count += len(text)

            # Zapisujemy to, co zostało po zakończeniu pętli (ostatnia strona)
            if current_page_text:
                joined_text = "\n".join(current_page_text).strip()
                if joined_text:
                    documents.append(Document(page_content=joined_text, metadata={"page_number": page_num}))

            return documents
        except Exception as e:
            logger.error(f"Błąd parsowania DOCX {file_path}: {e}")
            return []

    # --- Zwraca List[Document], każdy Document to jedna strona PDF ---
    def _process_pdf(self, file_path: str) -> List[Document]:
        """
        Prywatna metoda do ekstrakcji tekstu z pliku PDF z podziałem na strony.
        """
        documents = []
        try:
            reader = pypdf.PdfReader(file_path)

            for i, page in enumerate(reader.pages):
                extracted = page.extract_text()
                if extracted and extracted.strip():
                    # --- ZMIANA: Zapisujemy tekst strony do obiektu Document wraz z jej numerem ---
                    doc = Document(
                        page_content=extracted,
                        metadata={"page_number": i + 1}
                    )
                    documents.append(doc)

            return documents

        except Exception as e:
            logger.error(f"Błąd przetwarzania PDF {os.path.basename(file_path)}: {e}")
            return []

    # --- Zwraca List[Document] (jeden dokument dla plików tekstowych) ---
    def _process_text(self, file_path: str) -> List[Document]:
        """
        Odczyt plików tekstowych (TXT, MD, XML, JSON, XSD itp.).
        """
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
                if content.strip():
                    return [Document(page_content=content, metadata={"page_number": 1})]
                return []
        except Exception as e:
            logger.error(f"Błąd odczytu tekstu {os.path.basename(file_path)}: {e}")
            return []

    def _extract_domain_first(self, file_path: str) -> str:
        """
        Pobiera nazwę domeny na podstawie ścieżki pliku.
        Domena to pierwszy katalog względem głównego folderu przeszukiwania (self.directory).
        Np. dla pliku "C:/dane/HR/umowy/2023/umowa.pdf" (gdy self.directory="C:/dane")
        zwróci "HR".
        """
        # Wyliczamy ścieżkę względną pliku w stosunku do przeszukiwanego katalogu
        rel_path = os.path.relpath(file_path, self.directory)

        # Ujednolicamy ukośniki, aby split zachowywał się tak samo na Windowsie i Linuksie
        normalized_rel_path = rel_path.replace('\\', '/')
        parts = normalized_rel_path.split('/')

        if len(parts) > 1:
            # Plik leży w jakimś podfolderze. Bierzemy folder NAJWYŻSZEGO poziomu.
            return parts[0]
        else:
            # Plik leży bezpośrednio w głównym katalogu przeszukiwania.
            # Fallback na nazwę tego głównego katalogu lub "local"
            return os.path.basename(self.directory) or "local"